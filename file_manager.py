from phdictionary.requestor import Requester
from docx import Document
from random import randint


class FileManager:
    def txt_french_english(self, f, number_of_examples=0):
        r = Requester()
        f = f.read().splitlines()
        if number_of_examples == 0:
            document = Document()
            document.add_heading('French-English', 0)
            for word in f:
                if word != '':
                    re = r.get_french_english_q(word, 0)[:2]
                    if re[0] == 'feminine noun':
                        word += ' (f)'
                    elif re[0] == 'masculine noun':
                        word += ' (m) '
                    document.add_paragraph(word + '  -  ' + re[1], style='List Bullet')
            document.save(f'french-english{randint(0, 1000)}.docx')
        else:
            document = Document()
            document.add_heading('French-English', 0)
            for word in f:
                if word != '':
                    re = r.get_french_english_q(word, number_of_examples)
                    if re[0] == 'feminine noun':
                        word += ' (f)'
                    elif re[0] == 'masculine noun':
                        word += ' (m) '
                    p = document.add_paragraph()
                    p.add_run(word + '  -  ' + re[1]).bold = True
                    for ex in re[2]:
                        document.add_paragraph(ex, style='List Bullet')

            document.save(f'french-english-examples{randint(0, 1000)}.docx')

    def syn_txt(self, f, number_of_syn):
        r = Requester()
        f = f.read().splitlines()
        document = Document()
        document.add_heading('Synonyms', 0)
        for word in f:
            if word != '':
                re = r.get_synonym(str(word), number_of_syn)
                p = document.add_paragraph(style='List Bullet')
                to_add = ''

                for i in range(len(re)):
                    re[i].replace('%20', ' ')
                    if i == len(re) - 1:
                        to_add += re[i]
                    else:
                        to_add += re[i] + ', '
                p.add_run(word + '  -  ' + to_add).bold = True
        document.save(f'synonyms{randint(0, 1000)}.docx')

    def definition(self, f, number_of_examples):
        f = f.read().splitlines()
        r = Requester()
        if number_of_examples == 0:
            document = Document()
            document.add_heading('Definition', 0)
            for word in f:
                if word != '':
                    re = r.get_definition_1(word, 0)[:2]
                    if '\n' not in word:
                        p = document.add_paragraph(style='List Bullet')
                        p.add_run(word[0].upper() + word[1:]).bold = True
                        p.add_run('  -  ' + re[1])
                    else:
                        word = word[0].upper() + word[1:len(word) - 1]
                        p = document.add_paragraph(style='List Bullet')
                        p.add_run(word).bold = True
                        p.add_run('  -  ' + re[1])
            document.save(f'definition{randint(0, 1000)}.docx')
        else:
            document = Document()
            document.add_heading('Definition', 0)
            for word in f:
                if word != '':
                    re = r.get_definition_1(word, number_of_examples)
                    if '\n' not in word:
                        p = document.add_paragraph()
                        p.add_run(word[0].upper() + word[1:]).bold = True
                        p.add_run('  -  ' + re[1])
                    else:
                        word = word[0].upper() + word[1:len(word) - 1]
                        p = document.add_paragraph()
                        p.add_run(word).bold = True
                        p.add_run('  -  ' + re[1])
                    for ex in re[2]:
                        document.add_paragraph(ex, style='List Bullet')
            document.save(f'definition_examples{randint(0, 1000)}.docx')

    def txt_english_french(self, f, number_of_examples=0):
        r = Requester()
        f = f.read().splitlines()
        if number_of_examples == 0:
            document = Document()
            document.add_heading('English-French', 0)
            for word in f:
                if word != '':
                    re = r.get_eng_fr(word, 0)[:2]
                    if re[0] == 'feminine noun':
                        word += ' (f)'
                    elif re[0] == 'masculine noun':
                        word += ' (m) '
                    document.add_paragraph(word + '  -  ' + re[1], style='List Bullet')
            document.save(f'english-french{randint(0, 1000)}.docx')
        else:
            document = Document()
            document.add_heading('English-French', 0)
            for word in f:
                if word != '':
                    re = r.get_eng_fr(word, number_of_examples)
                    if re[0] == 'feminine noun':
                        word += ' (f)'
                    elif re[0] == 'masculine noun':
                        word += ' (m) '
                    p = document.add_paragraph()
                    p.add_run(word + '  -  ' + re[1]).bold = True
                    for ex in re[2]:
                        document.add_paragraph(ex, style='List Bullet')

            document.save(f'english-french-examples{randint(0, 1000)}.docx')
